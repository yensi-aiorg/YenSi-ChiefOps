"""
Google Drive folder processor.

Handles individual files uploaded from Google Drive exports. Detects
file types by extension, extracts text from supported formats (PDF,
DOCX, XLSX, MD, TXT, HTML), stores metadata in the ``drive_files``
collection, and prepares text documents for Citex RAG indexing.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from motor.motor_asyncio import AsyncIOMotorDatabase

from app.models.base import generate_uuid, utc_now
from app.services.ingestion.hasher import compute_hash, check_duplicate, record_hash
from app.services.ingestion.slack_admin import IngestionFileResult

logger = logging.getLogger(__name__)

# Supported text-extractable extensions
_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv"}
_HTML_EXTENSIONS = {".html", ".htm"}
_DOCX_EXTENSIONS = {".docx"}
_XLSX_EXTENSIONS = {".xlsx"}
_PDF_EXTENSIONS = {".pdf"}


async def process_drive_files(
    file_paths: list[str],
    db: AsyncIOMotorDatabase,  # type: ignore[type-arg]
) -> list[IngestionFileResult]:
    """Process a batch of files uploaded from Google Drive.

    Args:
        file_paths: List of file paths on disk.
        db: Motor database handle.

    Returns:
        A list of ``IngestionFileResult`` objects, one per file.
    """
    results: list[IngestionFileResult] = []

    for file_path in file_paths:
        result = await _process_single_file(file_path, db)
        results.append(result)

    logger.info(
        "Drive files processed: %d files, %d total records",
        len(results),
        sum(r.records_created for r in results),
    )
    return results


async def _process_single_file(
    file_path: str,
    db: AsyncIOMotorDatabase,  # type: ignore[type-arg]
) -> IngestionFileResult:
    """Process a single Drive file."""
    result = IngestionFileResult(filename=file_path, file_type="drive_document")

    if not os.path.isfile(file_path):
        result.status = "failed"
        result.errors.append(f"File does not exist: {file_path}")
        return result

    try:
        with open(file_path, "rb") as f:
            raw_content = f.read()
    except OSError as exc:
        result.status = "failed"
        result.errors.append(f"Cannot read file: {exc}")
        return result

    # Check for duplicates
    content_hash = compute_hash(raw_content)
    is_dup = await check_duplicate(content_hash, db)
    if is_dup:
        result.records_skipped += 1
        logger.debug("Skipping duplicate file: %s", file_path)
        return result

    filename = os.path.basename(file_path)
    ext = Path(file_path).suffix.lower()
    file_size = len(raw_content)

    # Extract text content based on file type
    extracted_text = await _extract_text(file_path, raw_content, ext, result)

    # Store file metadata in drive_files collection
    file_doc: dict[str, Any] = {
        "file_id": generate_uuid(),
        "filename": filename,
        "file_path": file_path,
        "extension": ext,
        "file_size": file_size,
        "file_hash": content_hash,
        "content_type": _get_content_type(ext),
        "has_extracted_text": extracted_text is not None and len(extracted_text) > 0,
        "text_length": len(extracted_text) if extracted_text else 0,
        "created_at": utc_now(),
        "updated_at": utc_now(),
    }

    await db.drive_files.insert_one(file_doc)
    await record_hash(content_hash, filename, db)
    result.records_created += 1

    # Build text document for Citex indexing if we got text
    if extracted_text:
        text_hash = compute_hash(extracted_text.encode("utf-8"))
        text_doc = {
            "document_id": generate_uuid(),
            "source": "gdrive",
            "source_ref": filename,
            "title": filename,
            "content": extracted_text,
            "content_hash": text_hash,
            "content_type": "text/plain",
            "original_format": ext,
            "file_size": file_size,
            "created_at": utc_now(),
        }
        await db.text_documents.update_one(
            {"source": "gdrive", "source_ref": filename},
            {"$set": text_doc},
            upsert=True,
        )

    return result


async def _extract_text(
    file_path: str,
    raw_content: bytes,
    ext: str,
    result: IngestionFileResult,
) -> Optional[str]:
    """Extract text content from a file based on its extension."""

    if ext in _TEXT_EXTENSIONS:
        return _extract_plain_text(raw_content)

    if ext in _HTML_EXTENSIONS:
        return _extract_html_text(raw_content)

    if ext in _DOCX_EXTENSIONS:
        return _extract_docx_text(file_path, result)

    if ext in _XLSX_EXTENSIONS:
        return _extract_xlsx_text(file_path, result)

    if ext in _PDF_EXTENSIONS:
        return _extract_pdf_text(raw_content, result)

    logger.debug("No text extraction available for extension: %s", ext)
    return None


def _extract_plain_text(content: bytes) -> str:
    """Decode plain text content."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return content.decode("utf-8", errors="replace")


def _extract_html_text(content: bytes) -> str:
    """Extract visible text from HTML by stripping tags."""
    text = _extract_plain_text(content)

    # Simple tag stripping without external dependencies
    import re

    # Remove script and style blocks
    text = re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    # Replace block-level tags with newlines
    text = re.sub(r"<(?:br|p|div|h[1-6]|li|tr)[^>]*>", "\n", text, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r"<[^>]+>", " ", text)
    # Decode common HTML entities
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = text.replace("&nbsp;", " ").replace("&quot;", '"')
    # Collapse whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _extract_docx_text(file_path: str, result: IngestionFileResult) -> Optional[str]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    except ImportError:
        result.errors.append("python-docx not installed, cannot extract DOCX text")
        return None
    except Exception as exc:
        result.errors.append(f"DOCX extraction error: {exc}")
        return None


def _extract_xlsx_text(file_path: str, result: IngestionFileResult) -> Optional[str]:
    """Extract text from an XLSX file using openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"## Sheet: {sheet_name}")
            lines.append("")

            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                # Skip entirely empty rows
                if any(c.strip() for c in cells):
                    lines.append(" | ".join(cells))

            lines.append("")

        wb.close()
        return "\n".join(lines)

    except ImportError:
        result.errors.append("openpyxl not installed, cannot extract XLSX text")
        return None
    except Exception as exc:
        result.errors.append(f"XLSX extraction error: {exc}")
        return None


def _extract_pdf_text(content: bytes, result: IngestionFileResult) -> Optional[str]:
    """Extract text from PDF using basic binary parsing.

    This implements a lightweight PDF text extraction that works without
    external system dependencies like poppler. It handles simple PDFs
    with plain text streams.
    """
    try:
        text = content.decode("latin-1", errors="replace")
        extracted_parts: list[str] = []

        # Find text between BT and ET markers (PDF text objects)
        import re

        # Extract text from Tj and TJ operators
        tj_pattern = re.compile(r"\(([^)]*)\)\s*Tj", re.DOTALL)
        for match in tj_pattern.finditer(text):
            part = match.group(1)
            # Unescape PDF string escapes
            part = part.replace("\\n", "\n").replace("\\r", "\r")
            part = part.replace("\\t", "\t").replace("\\(", "(").replace("\\)", ")")
            if part.strip():
                extracted_parts.append(part)

        # Also try TJ arrays
        tj_array_pattern = re.compile(r"\[([^\]]*)\]\s*TJ", re.DOTALL)
        for match in tj_array_pattern.finditer(text):
            array_content = match.group(1)
            # Extract strings from the array
            string_pattern = re.compile(r"\(([^)]*)\)")
            for s_match in string_pattern.finditer(array_content):
                part = s_match.group(1)
                part = part.replace("\\n", "\n").replace("\\r", "\r")
                part = part.replace("\\(", "(").replace("\\)", ")")
                if part.strip():
                    extracted_parts.append(part)

        if extracted_parts:
            return "\n".join(extracted_parts)

        # Fallback: no extractable text found
        result.errors.append("PDF text extraction found no text (may be image-based)")
        return None

    except Exception as exc:
        result.errors.append(f"PDF extraction error: {exc}")
        return None


def _get_content_type(ext: str) -> str:
    """Map file extension to MIME content type."""
    mime_map: dict[str, str] = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ".ppt": "application/vnd.ms-powerpoint",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".html": "text/html",
        ".htm": "text/html",
        ".csv": "text/csv",
        ".json": "application/json",
        ".xml": "application/xml",
        ".yaml": "application/yaml",
        ".yml": "application/yaml",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".svg": "image/svg+xml",
    }
    return mime_map.get(ext, "application/octet-stream")
